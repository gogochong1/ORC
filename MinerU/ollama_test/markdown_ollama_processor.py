import os
import re
import requests
from pathlib import Path
import base64
from PIL import Image
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class MarkdownOllamaProcessor:
    def __init__(self, ollama_api_url="http://localhost:11434/api/generate"):
        self.ollama_api_url = ollama_api_url
        self.model = "openstack:latest"
        self.embedding_model = "nomic-embed-text:latest"

    def read_markdown(self, file_path):
        """读取Markdown文件内容"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content

    def extract_questions_and_answers(self, content):
        """提取题目和答案"""
        # 提取选择题
        choice_pattern = r'(\d+)\.(.*?)\(([A-Z]+)\)(?=\s*\d+\.|\s*\d+．|$)'
        choices = re.findall(choice_pattern, content, re.DOTALL)
        
        # 提取判断题
        true_false_pattern = r'(\d+)．(.*?)\(([TF])\)(?=\s*\d+\.|\s*\d+．|$)'
        true_false = re.findall(true_false_pattern, content, re.DOTALL)
        
        # 打印调试信息
        print(f"找到 {len(choices)} 个选择题和 {len(true_false)} 个判断题")
        
        results = []
        
        # 处理选择题
        for num, question, answer in choices:
            question = question.strip()
            # 提取选项
            options = re.findall(r'([A-Z]\.\s*.*?)(?=[A-Z]\.|$)', question)
            # 保留原始题目文本，不删除选项
            question_text = question.strip()
            
            # 打印调试信息
            print(f"\n处理选择题 {num}:")
            print(f"问题内容: {question_text}")
            print(f"选项数量: {len(options)}")
            print(f"答案: {answer}")
            
            # 构建完整的题目信息
            full_question = {
                'type': 'choice',
                'number': num,
                'question': question_text,
                'options': options,
                'answer': answer,
                'analysis': None,
                'judgment': None,
                'reasoning': None,
                'suggestions': None
            }
            
            # 使用Ollama分析答案
            prompt = f"""作为Openstack专家，请分析以下题目：

题目：{question_text}
选项：
{chr(10).join(options)}
学生答案：{answer}

请按照以下格式提供分析：
1. 答案评估：[正确/错误]
2. 分析过程：
   - 题目理解
   - 选项分析
   - 答案验证
3. 知识点分析：
   - 考察的知识点
   - 相关概念
   - 技术要点
4. 改进建议：
   - 学习重点
   - 实践建议
   - 注意事项
5. 评分：[0-100分]

请基于Openstack知识库进行专业分析，并给出详细的评分理由。"""
            
            print(f"正在批改选择题 {num}...")
            analysis = self.process_with_ollama(prompt)
            if analysis:
                # 解析分析结果
                judgment = re.search(r'批改结果：\[(.*?)\]', analysis)
                reasoning = re.search(r'批改过程：\s*(.*?)(?=详细评析：|$)', analysis, re.DOTALL)
                detailed_analysis = re.search(r'详细评析：\s*(.*?)(?=改进建议：|$)', analysis, re.DOTALL)
                suggestions = re.search(r'改进建议：\s*(.*?)$', analysis, re.DOTALL)
                
                full_question['judgment'] = judgment.group(1) if judgment else "无法判断"
                full_question['reasoning'] = reasoning.group(1).strip() if reasoning else "无批改过程"
                full_question['analysis'] = detailed_analysis.group(1).strip() if detailed_analysis else "无详细评析"
                full_question['suggestions'] = suggestions.group(1).strip() if suggestions else "无改进建议"
            
            results.append(full_question)
        
        # 处理判断题
        for num, question, answer in true_false:
            question = question.strip()
            
            # 打印调试信息
            print(f"\n处理判断题 {num}:")
            print(f"问题内容: {question}")
            print(f"答案: {answer}")
            
            # 构建完整的题目信息
            full_question = {
                'type': 'true_false',
                'number': num,
                'question': question,
                'answer': answer,
                'analysis': None,
                'judgment': None,
                'reasoning': None,
                'suggestions': None
            }
            
            # 使用Ollama分析答案
            prompt = f"""作为Openstack专家，请分析以下判断题：

题目：{question}
学生答案：{answer}

请按照以下格式提供分析：
1. 答案评估：[正确/错误]
2. 分析过程：
   - 题目理解
   - 答案验证
3. 知识点分析：
   - 考察的知识点
   - 相关概念
   - 技术要点
4. 改进建议：
   - 学习重点
   - 实践建议
   - 注意事项
5. 评分：[0-100分]

请基于Openstack知识库进行专业分析，并给出详细的评分理由。"""
            
            print(f"正在批改判断题 {num}...")
            analysis = self.process_with_ollama(prompt)
            if analysis:
                # 解析分析结果
                judgment = re.search(r'批改结果：\[(.*?)\]', analysis)
                reasoning = re.search(r'批改过程：\s*(.*?)(?=详细评析：|$)', analysis, re.DOTALL)
                detailed_analysis = re.search(r'详细评析：\s*(.*?)(?=改进建议：|$)', analysis, re.DOTALL)
                suggestions = re.search(r'改进建议：\s*(.*?)$', analysis, re.DOTALL)
                
                full_question['judgment'] = judgment.group(1) if judgment else "无法判断"
                full_question['reasoning'] = reasoning.group(1).strip() if reasoning else "无批改过程"
                full_question['analysis'] = detailed_analysis.group(1).strip() if detailed_analysis else "无详细评析"
                full_question['suggestions'] = suggestions.group(1).strip() if suggestions else "无改进建议"
            
            results.append(full_question)
        
        return results

    def extract_images(self, markdown_content, base_path):
        """提取Markdown中的图片路径"""
        image_pattern = r'!\[.*?\]\((.*?)\)'
        images = re.findall(image_pattern, markdown_content)
        return [os.path.normpath(os.path.join(base_path, img)) for img in images]

    def encode_image(self, image_path):
        """将图片编码为base64"""
        try:
            print(f"正在处理图片: {image_path}")
            if not os.path.exists(image_path):
                print(f"图片文件不存在: {image_path}")
                return None
                
            with Image.open(image_path) as img:
                if img.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1])
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                buffer = io.BytesIO()
                img.save(buffer, format='JPEG')
                return base64.b64encode(buffer.getvalue()).decode('utf-8')
        except Exception as e:
            print(f"处理图片时出错 {image_path}: {str(e)}")
            return None

    def get_text_embedding(self, text):
        """获取文本的嵌入向量"""
        try:
            payload = {
                "model": self.embedding_model,
                "prompt": text,
                "options": {
                    "temperature": 0.1,
                    "top_p": 0.9
                }
            }
            
            response = requests.post(self.ollama_api_url, json=payload)
            response.raise_for_status()
            return response.json()["embedding"]
        except Exception as e:
            print(f"获取文本嵌入时出错: {str(e)}")
            return None

    def process_with_ollama(self, text, images=None):
        """使用Ollama处理文本和图片"""
        # 首先获取文本的嵌入向量
        text_embedding = self.get_text_embedding(text)
        
        messages = [{"role": "user", "content": text}]
        
        if images:
            for img_base64 in images:
                if img_base64:
                    messages.append({
                        "role": "user",
                        "content": [{"type": "image", "image": img_base64}]
                    })

        payload = {
            "model": self.model,
            "messages": messages,
            "stream": False,
            "options": {
                "temperature": 0.7,
                "top_p": 0.9,
                "num_ctx": 4096
            }
        }

        try:
            response = requests.post(self.ollama_api_url, json=payload)
            response.raise_for_status()
            return response.json()["response"]
        except Exception as e:
            print(f"调用Ollama API时出错: {str(e)}")
            return None

    def generate_doc(self, results, output_path):
        """生成Word文档"""
        try:
            # 检查文件是否已存在且被占用
            if os.path.exists(output_path):
                try:
                    # 尝试重命名现有文件
                    backup_path = f"{output_path}.bak"
                    if os.path.exists(backup_path):
                        os.remove(backup_path)
                    os.rename(output_path, backup_path)
                    print(f"已备份原文件到: {backup_path}")
                except PermissionError:
                    # 如果无法重命名，尝试使用新的文件名
                    base_name, ext = os.path.splitext(output_path)
                    counter = 1
                    while True:
                        new_path = f"{base_name}_{counter}{ext}"
                        if not os.path.exists(new_path):
                            output_path = new_path
                            print(f"使用新文件名: {output_path}")
                            break
                        counter += 1
            
            doc = Document()
            
            # 添加标题
            title = doc.add_paragraph('题目分析报告')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.size = Pt(16)
            title.runs[0].font.bold = True
            
            # 添加内容
            for item in results:
                # 添加题目编号和类型
                q_type = "选择题" if item['type'] == 'choice' else "判断题"
                q_para = doc.add_paragraph(f'第 {item["number"]} 题 ({q_type}):')
                q_para.runs[0].font.bold = True
                
                # 添加题目内容
                question_para = doc.add_paragraph()
                question_para.add_run(item['question']).font.size = Pt(12)
                
                # 如果是选择题，添加选项
                if item['type'] == 'choice':
                    for option in item['options']:
                        option_para = doc.add_paragraph()
                        option_para.add_run(option.strip()).font.size = Pt(11)
                
                # 添加答案
                a_para = doc.add_paragraph('原答案:')
                a_para.runs[0].font.bold = True
                doc.add_paragraph(item['answer'])
                
                # 添加模型处理过程
                process_para = doc.add_paragraph()
                process_para.add_run('模型处理过程：').font.bold = True
                doc.add_paragraph(f"使用模型 {self.model} 进行分析")
                
                # 添加判断结果
                judgment_para = doc.add_paragraph()
                judgment_para.add_run('判断结果：').font.bold = True
                judgment_para.add_run(item['judgment'] if item['judgment'] else "无法判断")
                
                # 添加判断依据
                reasoning_para = doc.add_paragraph()
                reasoning_para.add_run('判断依据：').font.bold = True
                if item['reasoning']:
                    for line in item['reasoning'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph("无判断依据")
                
                # 添加详细评析
                analysis_para = doc.add_paragraph()
                analysis_para.add_run('详细评析：').font.bold = True
                if item['analysis']:
                    for line in item['analysis'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph("无详细评析")
                
                # 添加改进建议
                suggestions_para = doc.add_paragraph()
                suggestions_para.add_run('改进建议：').font.bold = True
                if item['suggestions']:
                    for line in item['suggestions'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    doc.add_paragraph("无改进建议")
                
                # 添加分隔线
                doc.add_paragraph('-' * 50)
            
            # 保存文档
            doc.save(output_path)
            print(f"文档已成功保存到: {output_path}")
            
        except PermissionError as e:
            print(f"错误：无法保存文档。请检查文件权限或关闭可能正在使用该文件的程序。")
            print(f"错误详情: {str(e)}")
            # 尝试使用临时文件
            temp_path = "temp_analysis_results.docx"
            try:
                doc.save(temp_path)
                print(f"文档已保存到临时文件: {temp_path}")
            except Exception as e:
                print(f"无法保存到临时文件: {str(e)}")
        except Exception as e:
            print(f"保存文档时发生错误: {str(e)}")

def main():
    processor = MarkdownOllamaProcessor()
    
    # 读取Markdown文件
    markdown_path = "output/pdf_change/01.md"
    base_path = os.path.dirname(markdown_path)
    
    print(f"正在读取文件: {markdown_path}")
    content = processor.read_markdown(markdown_path)
    print(f"文件内容长度: {len(content)} 字符")
    
    questions = processor.extract_questions_and_answers(content)
    
    if not questions:
        print("警告：没有找到任何题目！")
        return
    
    # 生成文档
    output_doc = "analysis_results.docx"
    processor.generate_doc(questions, output_doc)
    print(f"\n分析报告已生成：{output_doc}")
    print(f"共处理了 {len(questions)} 个题目")

if __name__ == "__main__":
    main()