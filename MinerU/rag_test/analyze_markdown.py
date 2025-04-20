import os
import re
import requests
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import base64
import io
import json

class MarkdownAnalyzer:
    def __init__(self):
        self.ollama_api_url = "http://localhost:11434/api/generate"
        self.vision_model = "llama3.2-vision:latest"
        self.embedding_model = "nomic-embed-text:latest"
        self.manage_knowledge = "Openstack"

    def read_markdown(self, file_path):
        """读取Markdown文件内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"成功读取文件: {file_path}")
            print(f"文件内容长度: {len(content)} 字符")
            return content
        except Exception as e:
            print(f"读取文件时出错: {str(e)}")
            return None

    def extract_images(self, content, base_path):
        """提取Markdown中的图片路径"""
        image_pattern = r'!\[.*?\]\((.*?)\)'
        images = re.findall(image_pattern, content)
        return [os.path.normpath(os.path.join(base_path, img)) for img in images]

    def encode_image(self, image_path):
        """将图片编码为base64"""
        try:
            if not os.path.exists(image_path):
                print(f"图片文件不存在: {image_path}")
                return None
                
            with Image.open(image_path) as img:
                if img.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1])
                    img = background
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                buffer = io.BytesIO()
                img.save(buffer, format='JPEG')
                return base64.b64encode(buffer.getvalue()).decode('utf-8')
        except Exception as e:
            print(f"处理图片时出错 {image_path}: {str(e)}")
            return None

    def extract_questions_and_answers(self, content):
        """使用Ollama模型智能提取Markdown格式的题目和答案"""
        if not content:
            print("错误：内容为空")
            return []
            
        results = []
        print("开始提取题目...")
        
        try:
            # 构建提示词
            prompt = f"""
            请分析以下Markdown文档内容，识别并提取所有题目。文档内容如下：

            {content}

            请直接返回处理后的Markdown格式内容，要求：
            1. 保持原有的Markdown格式
            2. 每个题目单独成段
            3. 题目格式为：
               ### 题号. 题目内容
               答案：xxx
               选项：
               A. xxx
               B. xxx
               ...
            4. 对于填空题，用___表示填空位置
            5. 对于简答题，保留完整的题目内容

            请确保：
            1. 准确识别题目编号
            2. 正确分类题型
            3. 提取完整的题目内容
            4. 保留所有选项（如果是选择题）
            5. 保持原有的格式和结构
            """
            
            # 调用Ollama API
            payload = {
                "model": self.vision_model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False
            }
            
            response = requests.post(self.ollama_api_url, json=payload)
            response.raise_for_status()
            processed_content = response.json()["response"]
            
            # 处理返回的Markdown内容
            current_question = None
            current_type = None
            current_options = []
            
            for line in processed_content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                    
                # 识别题目
                if line.startswith('###'):
                    # 保存上一个题目
                    if current_question:
                        results.append({
                            'number': current_question['number'],
                            'type': current_type,
                            'question': current_question['content'],
                            'answer': current_question.get('answer'),
                            'options': current_options.copy()
                        })
                    
                    # 开始新题目
                    match = re.match(r'###\s*(\d+)\.\s*(.*)', line)
                    if match:
                        number, content = match.groups()
                        current_question = {
                            'number': number,
                            'content': content
                        }
                        current_options = []
                        # 根据内容判断题型
                        if '___' in content:
                            current_type = 'fill'
                        elif any(opt in content for opt in ['A.', 'B.', 'C.', 'D.']):
                            current_type = 'choice'
                        else:
                            current_type = 'essay'
                
                # 处理答案
                elif line.startswith('答案：'):
                    if current_question:
                        current_question['answer'] = line[3:].strip()
                
                # 处理选项
                elif re.match(r'[A-Z]\.\s', line):
                    if current_type == 'choice':
                        current_options.append(line[2:].strip())
            
            # 保存最后一个题目
            if current_question:
                results.append({
                    'number': current_question['number'],
                    'type': current_type,
                    'question': current_question['content'],
                    'answer': current_question.get('answer'),
                    'options': current_options.copy()
                })
            
            print(f"总共提取了 {len(results)} 个题目")
            return results
            
        except Exception as e:
            print(f"提取题目时出错: {str(e)}")
            return []

    def analyze_question(self, question, images=None):
        """使用llama3.2-vision分析题目"""
        try:
            print(f"开始分析第 {question['number']} 题...")
            
            # 构建提示词
            prompt = f"""
            请分析以下题目：
            题目类型：{question['type']}
            题目内容：{question['question']}
            学生答案：{question['answer']}
            
            请按照以下格式回答：
            批改结果：[正确/错误/部分正确]
            批改过程：
            详细评析：
            改进建议：
            """
            
            messages = [{"role": "user", "content": prompt}]
            
            # 如果有图片，添加到消息中
            if images:
                for img_base64 in images:
                    if img_base64:
                        messages.append({
                            "role": "user",
                            "content": [{"type": "image", "image": img_base64}]
                        })
            
            payload = {
                "model": self.vision_model,
                "messages": messages,
                "stream": False
            }
            
            response = requests.post(self.ollama_api_url, json=payload)
            response.raise_for_status()
            analysis = response.json()["response"]
            
            # 解析分析结果
            judgment = re.search(r'批改结果：\[(.*?)\]', analysis)
            reasoning = re.search(r'批改过程：\s*(.*?)(?=详细评析：|$)', analysis, re.DOTALL)
            detailed_analysis = re.search(r'详细评析：\s*(.*?)(?=改进建议：|$)', analysis, re.DOTALL)
            suggestions = re.search(r'改进建议：\s*(.*?)$', analysis, re.DOTALL)
            
            result = {
                'judgment': judgment.group(1) if judgment else "无法判断",
                'reasoning': reasoning.group(1).strip() if reasoning else "无批改过程",
                'analysis': detailed_analysis.group(1).strip() if detailed_analysis else "无详细评析",
                'suggestions': suggestions.group(1).strip() if suggestions else "无改进建议"
            }
            
            print(f"第 {question['number']} 题分析完成")
            return result
            
        except Exception as e:
            print(f"分析题目时出错: {str(e)}")
            return None

    def generate_docx(self, questions, output_path):
        """生成Word文档报告"""
        try:
            print(f"开始生成文档: {output_path}")
            doc = Document()
            
            # 添加标题
            title = doc.add_paragraph('题目分析报告')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.size = Pt(16)
            title.runs[0].font.bold = True
            
            # 添加内容
            for question in questions:
                print(f"正在处理第 {question['number']} 题...")
                
                # 添加题目编号和类型
                q_type = {
                    'choice': '选择题',
                    'fill': '填空题',
                    'essay': '简答题'
                }.get(question['type'], '未知题型')
                
                q_para = doc.add_paragraph(f'第 {question["number"]} 题 ({q_type}):')
                q_para.runs[0].font.bold = True
                
                # 添加题目内容
                question_para = doc.add_paragraph()
                question_para.add_run(question['question']).font.size = Pt(12)
                
                # 如果是选择题，添加选项
                if question['type'] == 'choice':
                    for option in question['options']:
                        option_para = doc.add_paragraph()
                        option_para.add_run(option.strip()).font.size = Pt(11)
                
                # 添加答案
                a_para = doc.add_paragraph('原答案:')
                a_para.runs[0].font.bold = True
                if isinstance(question['answer'], list):
                    doc.add_paragraph('\n'.join(question['answer']))
                else:
                    doc.add_paragraph(str(question['answer']))
                
                # 添加分析结果
                if 'analysis_result' in question:
                    result = question['analysis_result']
                    
                    # 添加判断结果
                    judgment_para = doc.add_paragraph()
                    judgment_para.add_run('判断结果：').font.bold = True
                    judgment_para.add_run(result['judgment'])
                    
                    # 添加判断依据
                    reasoning_para = doc.add_paragraph()
                    reasoning_para.add_run('判断依据：').font.bold = True
                    if result['reasoning']:
                        for line in result['reasoning'].split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style='List Bullet')
                    else:
                        doc.add_paragraph("无判断依据")
                    
                    # 添加详细评析
                    analysis_para = doc.add_paragraph()
                    analysis_para.add_run('详细评析：').font.bold = True
                    if result['analysis']:
                        for line in result['analysis'].split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style='List Bullet')
                    else:
                        doc.add_paragraph("无详细评析")
                    
                    # 添加改进建议
                    suggestions_para = doc.add_paragraph()
                    suggestions_para.add_run('改进建议：').font.bold = True
                    if result['suggestions']:
                        for line in result['suggestions'].split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip(), style='List Bullet')
                    else:
                        doc.add_paragraph("无改进建议")
                
                # 添加分隔线
                doc.add_paragraph('-' * 50)
            
            # 保存文档
            doc.save(output_path)
            print(f"文档已成功保存到: {output_path}")
            
        except Exception as e:
            print(f"生成文档时出错: {str(e)}")

def main():
    analyzer = MarkdownAnalyzer()
    
    # 设置输入输出目录
    input_dir = "rag_test/pdf_output"
    output_dir = "rag_test/md_output"
    
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 遍历输入目录中的所有markdown文件
    for filename in os.listdir(input_dir):
        if filename.lower().endswith('.md'):
            md_path = os.path.join(input_dir, filename)
            print(f"\n开始处理文件: {filename}")
            
            try:
                # 读取markdown文件
                content = analyzer.read_markdown(md_path)
                if not content:
                    print(f"错误：无法读取文件 {filename}")
                    continue
                
                # 提取图片
                base_path = os.path.dirname(md_path)
                images = analyzer.extract_images(content, base_path)
                encoded_images = [analyzer.encode_image(img) for img in images]
                print(f"找到 {len(images)} 张图片")
                
                # 提取题目和答案
                questions = analyzer.extract_questions_and_answers(content)
                
                if not questions:
                    print(f"警告：在文件 {filename} 中没有找到任何题目！")
                    continue
                
                # 分析每个题目
                for question in questions:
                    analysis_result = analyzer.analyze_question(question, encoded_images)
                    if analysis_result:
                        question['analysis_result'] = analysis_result
                
                # 生成报告
                output_docx = os.path.join(output_dir, f"{os.path.splitext(filename)[0]}_analysis.docx")
                analyzer.generate_docx(questions, output_docx)
                print(f"成功处理文件: {filename}")
                
            except Exception as e:
                print(f"处理文件 {filename} 时出错: {str(e)}")

if __name__ == "__main__":
    main() 