from student_rag import StudentRag
from kb import Kb
import os
from pathlib import Path
import docx

def create_knowledge_base(input_dir, output_file):
    """从输入目录创建知识库文件"""
    print(f"正在从目录 {input_dir} 创建知识库...")
    
    with open(output_file, 'w', encoding='utf-8') as outfile:
        # 遍历目录中的所有文件
        for filename in os.listdir(input_dir):
            filepath = os.path.join(input_dir, filename)
            if not os.path.isfile(filepath):
                continue
                
            print(f"处理文件: {filename}")
            
            # 根据文件类型提取文本
            try:
                if filename.lower().endswith('.pdf'):
                    content = Kb.extract_text_from_pdf(filepath)
                elif filename.lower().endswith('.docx'):
                    content = Kb.extract_text_from_docx(filepath)
                else:
                    print(f"跳过不支持的文件类型: {filename}")
                    continue
                
                # 写入文件名作为标题
                outfile.write(f"\n# {filename}\n\n")
                outfile.write(content)
                outfile.write("\n\n")
                
                print(f"已处理文件: {filename}")
                
            except Exception as e:
                print(f"处理文件 {filename} 时出错: {str(e)}")
                continue

def process_markdown_files(input_dir, output_dir, rag):
    """处理输入目录中的所有markdown文件"""
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    
    # 确保输出目录存在
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 遍历所有markdown文件
    for md_file in input_path.glob('*.md'):
        print(f"处理文件: {md_file}")
        
        # 读取markdown文件内容
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建一个新的docx文档
        doc = docx.Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal'].font.size = docx.shared.Pt(12)
        
        # 添加文件标题
        doc.add_heading(f"{md_file.stem} 批改结果", level=1)
        
        # 按题目类型分割内容
        sections = content.split('# ')
        for section in sections:
            if not section.strip():
                continue
                
            # 处理选择题
            if '选择题' in section:
                doc.add_heading('一、选择题', level=2)
                questions = section.split('\n\n')
                for q in questions:
                    if not q.strip() or q.strip().isdigit():
                        continue
                    # 提取题目和答案
                    lines = q.split('\n')
                    question = lines[0].strip()
                    answer = None
                    for line in lines[1:]:
                        if line.strip().startswith(('A、', 'B、', 'C、', 'D、')):
                            if '（' in line and '）' in line:
                                answer = line[line.find('（')+1:line.find('）')].strip()
                                break
                    if answer:
                        correction_result = rag.process_question(question, answer)
                        # 添加批改结果到文档
                        doc.add_paragraph(correction_result)
            
            # 处理填空题
            elif '填空题' in section:
                doc.add_heading('二、填空题', level=2)
                questions = section.split('\n\n')
                for q in questions:
                    if not q.strip() or q.strip().isdigit():
                        continue
                    # 提取题目和答案
                    question = q.strip()
                    answer = None
                    if '___' in question:
                        answer = question[question.find('___')+3:].strip()
                        question = question[:question.find('___')].strip()
                    if answer:
                        correction_result = rag.process_question(question, answer)
                        # 添加批改结果到文档
                        doc.add_paragraph(correction_result)
            
            # 处理简答题
            elif '简答题' in section:
                doc.add_heading('三、简答题', level=2)
                questions = section.split('\n\n')
                for q in questions:
                    if not q.strip() or q.strip().isdigit():
                        continue
                    # 提取题目和答案
                    lines = q.split('\n')
                    question = lines[0].strip()
                    answer = '\n'.join(lines[1:]).strip() if len(lines) > 1 else ''
                    if answer:
                        correction_result = rag.process_question(question, answer)
                        # 添加批改结果到文档
                        doc.add_paragraph(correction_result)
        
        # 保存文档
        output_file = output_path / f"{md_file.stem}_批改结果.docx"
        doc.save(str(output_file))
        print(f"已生成批改结果文件: {output_file}")

def main():
    # 获取当前脚本所在目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = os.path.dirname(current_dir)  # 上一级目录
    
    # 配置参数
    model = "qwq:latest"  # 使用qwq模型
    kb_filepath = os.path.join(current_dir, "knowledge_base.txt")  # 知识库文件路径
    pdf_dir = os.path.join(base_dir, "rag_kj", "pdf_01")  # PDF文件目录
    pdf_output_dir = os.path.join(base_dir, "pdf_output")  # PDF输出目录
    md_output_dir = os.path.join(base_dir, "md_output")    # Markdown输出目录
    
    # 创建知识库文件
    create_knowledge_base(pdf_dir, kb_filepath)
    print(f"知识库文件已创建: {kb_filepath}")
    
    # 初始化RAG系统
    rag = StudentRag(model, kb_filepath)
    
    # 处理所有markdown文件
    process_markdown_files(pdf_output_dir, md_output_dir, rag)

if __name__ == "__main__":
    main() 