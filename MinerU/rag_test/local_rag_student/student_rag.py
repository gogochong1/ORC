from kb import Kb
from ollama import Client
import re
import docx
from pathlib import Path
import PyPDF2

class StudentRag:
    def __init__(self, model, kb_filepath):
        self.kb_filepath = kb_filepath
        self.kb = Kb(kb_filepath)
        self.model = model
        self.client = Client(host='http://localhost:11434')
        self.question_types = {
            'choice': r'[A-D]、',  # 选择题特征
            'fill': r'_{2,}',      # 填空题特征（两个以上的下划线）
        }

    def identify_question_type(self, question):
        """识别题目类型"""
        if re.search(self.question_types['choice'], question):
            return 'choice'
        elif re.search(self.question_types['fill'], question):
            return 'fill'
        else:
            return 'essay'  # 默认为简答题

    def extract_choice_options(self, question):
        """提取选择题选项"""
        options = {}
        for line in question.split('\n'):
            match = re.search(r'([A-D])、(.*)', line)
            if match:
                options[match.group(1)] = match.group(2).strip()
        return options

    def extract_student_answer(self, question):
        """提取学生答案"""
        # 选择题答案提取
        choice_match = re.search(r'（\s*([A-D])\s*）', question)
        if choice_match:
            return choice_match.group(1)
        
        # 填空题答案提取
        fill_match = re.search(r'答案[:：]\s*(.*?)(?=\n|$)', question)
        if fill_match:
            return fill_match.group(1).strip()
        
        # 简答题答案提取（假设简答题答案在问题后面）
        essay_match = re.search(r'\n\n(.*?)(?=\n\n|$)', question, re.DOTALL)
        if essay_match:
            return essay_match.group(1).strip()
        
        return None

    def process_question(self, question, student_answer):
        """处理题目并生成批改结果"""
        question_type = self.identify_question_type(question)
        
        # 获取知识库上下文
        context = self.kb.search(question)
        
        # 构建提示词
        if question_type == 'choice':
            # 提取选项
            options = self.extract_choice_options(question)
            options_text = '\n'.join([f"{k}、{v}" for k, v in options.items()])
            
            system_prompt = """你是一个大学老师，你的职责是修改一份学生的作业，你需要判断他的回答是否正确，如果他的答案是错误的，请给出正确答案并且给出理由。

请按照以下格式输出：

题目：
[完整题目文本]

选项：
[所有选项，每行一个]

学生作答：[学生选择的选项]

评阅意见：
[详细分析学生答案的正确性，包括关键概念理解和技术要点，解释为什么这个答案是正确的或错误的]

正确答案：[正确的选项]
得分情况：[正确/错误]
--------------------------------------------------"""
            
            prompt = f"""请基于以下知识库内容，以OpenStack课程教师的身份评阅这道题目：

知识库参考内容：
{context}

待批改的题目：
{question}

选项：
{options_text}

学生答案：
{student_answer}

请严格按照系统提示的格式输出评阅结果，注意评分的公平性和评语的专业性。评阅意见要简洁明了，控制在200字以内。不要输出任何<think></think>标签中的内容。"""
            
        elif question_type == 'fill':
            system_prompt = """你是一个大学老师，你的职责是修改一份学生的作业，你需要判断他的回答是否正确，如果他的答案是错误的，请给出正确答案并且给出理由。

请按照以下格式输出：

题目：
[完整题目文本]

学生作答：[学生填写的答案]

评阅意见：
[详细分析学生答案的准确性，包括专业术语使用是否恰当，解释为什么这个答案是正确的或错误的]

正确答案：[标准答案]
得分情况：[正确/错误]
--------------------------------------------------"""
            
            prompt = f"""请基于以下知识库内容，以OpenStack课程教师的身份评阅这道题目：

知识库参考内容：
{context}

待批改的题目：
{question}

学生答案：
{student_answer}

请严格按照系统提示的格式输出评阅结果，注意评分的公平性和评语的专业性。评阅意见要简洁明了，控制在200字以内。不要输出任何<think></think>标签中的内容。"""
            
        else:  # 简答题
            system_prompt = """你是一个大学老师，你的职责是修改一份学生的作业，你需要判断他的回答是否正确，如果他的答案是错误的，请给出正确答案并且给出理由。

请按照以下格式输出：

题目：
[完整题目文本]

学生作答：
[学生的答案全文]

评阅意见：
[详细分析学生答案的完整性、专业性和逻辑性，从多个维度评价答案质量]

正确答案：
[完整的参考答案]

得分情况：[正确/错误]
--------------------------------------------------"""
            
            prompt = f"""请基于以下知识库内容，以OpenStack课程教师的身份评阅这道题目：

知识库参考内容：
{context}

待批改的题目：
{question}

学生答案：
{student_answer}

请严格按照系统提示的格式输出评阅结果，注意评分的公平性和评语的专业性。评阅意见要简洁明了，控制在200字以内。不要输出任何<think></think>标签中的内容。"""
        
        # 调用模型获取回答
        response = self.client.chat(model=self.model, messages=[
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': prompt}
        ])
        
        return response['message']['content']

    def process_exam(self, exam_content):
        """处理整份试卷"""
        # 分割不同类型的题目
        sections = exam_content.split('#')
        results = []
        
        for section in sections:
            if not section.strip():
                continue
                
            # 分割题目
            questions = section.strip().split('\n\n')
            section_title = questions[0].strip()
            results.append(f"# {section_title}\n")
            
            for question in questions[1:]:
                if not question.strip():
                    continue
                    
                # 处理每道题目
                result = self.process_question(question.strip())
                results.append(result)
                results.append('\n')
        
        return '\n'.join(results)

    def save_to_docx(self, question, student_answer, correction_result, output_path):
        """将题目、学生答案和批改结果保存为docx文件"""
        doc = docx.Document()
        
        # 设置字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal'].font.size = docx.shared.Pt(12)
        
        # 移除<think>标签及其内容
        correction_result = re.sub(r'<think>.*?</think>', '', correction_result, flags=re.DOTALL)
        
        # 解析批改结果
        correction_parts = correction_result.split('\n')
        
        for line in correction_parts:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('题目：'):
                doc.add_paragraph('题目：')
                doc.add_paragraph(line[3:].strip())
            elif line.startswith('选项：'):
                doc.add_paragraph('选项：')
                doc.add_paragraph(line[3:].strip())
            elif line.startswith('学生作答：'):
                doc.add_paragraph('学生作答：')
                doc.add_paragraph(line[5:].strip())
            elif line.startswith('评阅意见：'):
                doc.add_paragraph('评阅意见：')
                doc.add_paragraph(line[5:].strip())
            elif line.startswith('正确答案：'):
                doc.add_paragraph(line)
            elif line.startswith('得分情况：'):
                doc.add_paragraph(line)
        
        # 添加分隔线
        doc.add_paragraph('--------------------------------------------------')
        
        # 确保输出目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # 保存文档
        doc.save(output_path)

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """从PDF文件中提取文本"""
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text 