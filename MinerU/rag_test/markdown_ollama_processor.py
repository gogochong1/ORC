import os
import re
import requests
from pathlib import Path
import base64
from PIL import Image
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

class MarkdownOllamaProcessor:
    def __init__(self, ollama_api_url="http://localhost:8080/api/generate"):
        self.ollama_api_url = "http://localhost:8080/api/generate"
        self.ollama_chat_url = "http://localhost:8080/api/chat"
        self.ollama_embeddings_url = "http://localhost:8080/api/embeddings"
        self.vision_model = "qwq:latest"
        self.embed_model = "nomic-embed-text:latest"
        self.knowledge_base = {}  # 存储知识库的嵌入向量
        self.check_models()
        self.load_knowledge_base()

    def check_models(self):
        """检查模型是否已加载"""
        try:
            # 检查qwq模型
            chat_payload = {
                "model": self.vision_model,
                "messages": [
                    {
                        "role": "system",
                        "content": "系统初始化测试"
                    },
                    {
                        "role": "user",
                        "content": "测试消息"
                    }
                ]
            }
            
            print(f"正在检查 {self.vision_model} 模型...")
            headers = {
                "Content-Type": "application/json"
            }
            response = requests.post(self.ollama_chat_url, json=chat_payload, headers=headers, timeout=120)
            if response.status_code == 404:
                print(f"错误：找不到模型 {self.vision_model}，请确保已经安装")
                raise Exception(f"模型 {self.vision_model} 未找到")
            response.raise_for_status()
            print(f"模型 {self.vision_model} 检查通过")
            
            # 检查embedding模型
            embed_payload = {
                "model": self.embed_model,
                "prompt": "测试文本"
            }
            
            print(f"正在检查 {self.embed_model} 模型...")
            response = requests.post(self.ollama_embeddings_url, json=embed_payload, headers=headers, timeout=120)
            if response.status_code == 404:
                print(f"错误：找不到模型 {self.embed_model}，请确保已经安装")
                raise Exception(f"模型 {self.embed_model} 未找到")
            response.raise_for_status()
            print(f"模型 {self.embed_model} 检查通过")
            
        except requests.exceptions.RequestException as e:
            print(f"连接Open WebUI服务失败: {str(e)}")
            print("请确保Open WebUI服务正在运行，并且已经安装了所需的模型")
            print("可以使用以下命令安装模型：")
            print(f"ollama pull {self.vision_model}")
            print(f"ollama pull {self.embed_model}")
            raise
        except Exception as e:
            print(f"检查模型时出错: {str(e)}")
            raise

    def wait_for_model(self, model_name, max_attempts=30, delay=10):
        """等待模型加载完成"""
        for attempt in range(max_attempts):
            try:
                if "embed" in model_name:
                    payload = {
                        "model": model_name,
                        "prompt": "test"
                    }
                    response = requests.post("http://localhost:11434/api/embeddings", json=payload, timeout=120)
                else:
                    payload = {
                        "model": model_name,
                        "prompt": "test",
                        "stream": False
                    }
                    response = requests.post(self.ollama_api_url, json=payload, timeout=120)
                
                if response.status_code != 500 or "loading model" not in response.text:
                    print(f"模型 {model_name} 加载完成")
                    return True
                
                print(f"等待模型 {model_name} 加载... ({attempt + 1}/{max_attempts})")
                time.sleep(delay)
                
            except Exception as e:
                print(f"检查模型状态时出错: {str(e)}")
                time.sleep(delay)
                continue
        
        print(f"等待模型 {model_name} 加载超时")
        return False

    def load_knowledge_base(self):
        """加载Openstack知识库内容并生成嵌入向量"""
        try:
            # 从rag_kj目录加载Openstack知识库文件
            knowledge_dir = "rag_kj"
            if not os.path.exists(knowledge_dir):
                print(f"警告：Openstack知识库目录 {knowledge_dir} 不存在")
                return

            print("正在加载Openstack知识库...")
            for filename in os.listdir(knowledge_dir):
                if filename.endswith(('.txt', '.md')):
                    file_path = os.path.join(knowledge_dir, filename)
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                            # 将内容分成较小的块
                            chunks = self.split_into_chunks(content)
                            print(f"处理文件 {filename}，分割为 {len(chunks)} 个文本块")
                            # 为每个块生成嵌入向量
                            for chunk in chunks:
                                embedding = self.get_embedding(chunk)
                                if embedding:
                                    self.knowledge_base[chunk] = embedding
                    except Exception as e:
                        print(f"处理文件 {filename} 时出错: {str(e)}")
                        continue

            print(f"Openstack知识库加载完成，共处理了 {len(self.knowledge_base)} 个文本块")
        except Exception as e:
            print(f"加载Openstack知识库时出错: {str(e)}")

    def split_into_chunks(self, text, chunk_size=1000):
        """将文本分割成较小的块"""
        sentences = re.split(r'[。！？.!?]', text)
        chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if current_length + len(sentence) > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_length = 0

            current_chunk.append(sentence)
            current_length += len(sentence)

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks

    def get_embedding(self, text):
        """使用nomic-embed-text模型获取文本的嵌入向量"""
        try:
            # 确保文本不为空且是字符串类型
            if not text or not isinstance(text, str):
                print("警告：无效的文本输入")
                return None

            text = text.strip()
            if not text:
                print("警告：文本为空")
                return None

            payload = {
                "model": self.embed_model,
                "prompt": text
            }

            headers = {
                "Content-Type": "application/json"
            }

            response = requests.post(self.ollama_embeddings_url, json=payload, headers=headers)
            response.raise_for_status()
            result = response.json()
            
            if "error" in result:
                print(f"获取嵌入向量时出错: {result['error']}")
                return None
                
            if "embedding" not in result:
                print("API返回的响应中没有找到'embedding'字段")
                return None
                
            return result["embedding"]
        except requests.exceptions.RequestException as e:
            print(f"获取嵌入向量时出错: {str(e)}")
            return None
        except Exception as e:
            print(f"处理嵌入向量时出错: {str(e)}")
            return None

    def find_relevant_knowledge(self, query, top_k=3):
        """查找与查询最相关的知识片段"""
        try:
            if not self.knowledge_base:
                print("警告：知识库为空")
                return []

            print(f"正在为查询生成嵌入向量：{query}")
            query_embedding = self.get_embedding(query)
            if not query_embedding:
                print("警告：无法为查询生成嵌入向量")
                return []

            # 计算与所有知识块的相似度
            similarities = []
            for text, embedding in self.knowledge_base.items():
                similarity = self.cosine_similarity(query_embedding, embedding)
                similarities.append((text, similarity))

            # 按相似度排序并返回前top_k个结果
            similarities.sort(key=lambda x: x[1], reverse=True)
            return [text for text, sim in similarities[:top_k] if sim > 0.5]  # 只返回相似度大于0.5的结果
        except Exception as e:
            print(f"查找相关知识时出错: {str(e)}")
            return []

    def cosine_similarity(self, vec1, vec2):
        """计算两个向量的余弦相似度"""
        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        norm1 = sum(a * a for a in vec1) ** 0.5
        norm2 = sum(b * b for b in vec2) ** 0.5
        return dot_product / (norm1 * norm2) if norm1 * norm2 != 0 else 0

    def process_with_ollama(self, text, images=None, max_retries=3):
        """使用Ollama处理文本和图片，结合知识库"""
        for attempt in range(max_retries):
            try:
                # 使用nomic-embed-text查找相关知识
                print("正在使用nomic-embed-text查找相关知识...")
                relevant_knowledge = self.find_relevant_knowledge(text)
                
                if not relevant_knowledge:
                    print("警告：在知识库中未找到相关内容")
                    knowledge_context = "未在知识库中找到相关内容。"
                else:
                    knowledge_context = "\n".join(relevant_knowledge)
                    print(f"找到 {len(relevant_knowledge)} 条相关知识")

                # 构建消息
                messages = [
                    {
                        "role": "system",
                        "content": "你是一个专业的助手，请基于提供的知识库内容回答问题。回答必须包含'推理过程'和'答案'两个部分。"
                    },
                    {
                        "role": "user",
                        "content": f"""基于以下知识：

{knowledge_context}

请回答问题：{text}

要求：
1. 必须基于上述知识进行回答
2. 回答必须包含推理过程和最终答案
3. 如果知识库中没有相关信息，请明确指出"""
                    }
                ]

                # 使用qwq模型生成回答
                print("正在生成回答...")
                chat_payload = {
                    "model": self.vision_model,
                    "messages": messages
                }

                headers = {
                    "Content-Type": "application/json"
                }

                response = requests.post(self.ollama_chat_url, json=chat_payload, headers=headers, timeout=120)
                response.raise_for_status()
                result = response.json()
                
                if "error" in result:
                    raise Exception(f"生成回答时出错: {result['error']}")
                
                if "message" not in result:
                    raise Exception("返回的响应中没有找到'message'字段")
                
                result_text = result["message"]["content"]
                
                # 如果返回的内容没有包含必要的格式，添加格式化处理
                if "推理过程：" not in result_text or "答案：" not in result_text:
                    # 将返回的内容格式化
                    formatted_response = f"""推理过程：
基于知识库内容分析：
{result_text}

答案：
基于上述推理，答案是：{result_text}"""
                    result_text = formatted_response

                # 从结果中提取推理过程和答案
                reasoning = re.search(r'推理过程：\s*(.*?)(?=答案：|$)', result_text, re.DOTALL)
                reasoning_text = reasoning.group(1).strip() if reasoning else "无推理过程"

                answer_match = re.search(r'答案：\s*(.*?)(?=推理过程：|$)', result_text, re.DOTALL)
                if not answer_match:
                    answer_match = re.search(r'答案：\s*\[(.*?)\]', result_text)

                inferred_answer = answer_match.group(1).strip() if answer_match else "无法确定"

                # 验证结果
                if reasoning_text == "无推理过程" or inferred_answer == "无法确定":
                    if attempt < max_retries - 1:
                        print("响应格式不正确，重试中...")
                        continue

                print("回答生成完成")
                return {
                    'reasoning': reasoning_text,
                    'answer': inferred_answer,
                    'knowledge_context': knowledge_context
                }

            except requests.exceptions.RequestException as e:
                print(f"API请求错误 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
                if attempt == max_retries - 1:
                    print("达到最大重试次数，放弃处理")
                    return None
                time.sleep(10)
                continue
            except Exception as e:
                print(f"处理时出错 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
                if attempt == max_retries - 1:
                    print("达到最大重试次数，放弃处理")
                    return None
                time.sleep(10)
                continue

    def extract_questions_and_answers(self, content):
        """提取题目和答案"""
        # 使用数字加顿号作为分割符
        question_pattern = r'(\d+、.*?)(?=\d+、|$)'
        questions = re.findall(question_pattern, content, re.DOTALL)
        
        # 打印调试信息
        print(f"找到 {len(questions)} 个题目")
        
        results = []
        
        # 处理每个题目
        for question in questions:
            try:
                # 提取题号
                num_match = re.match(r'(\d+)、', question)
                if not num_match:
                    continue
                    
                num = num_match.group(1)
                question_text = question[len(num_match.group(0)):].strip()
                
                # 判断题目类型
                if '___' in question_text:
                    question_type = 'fill_in'
                elif '？' in question_text or '?' in question_text:
                    question_type = 'essay'
                else:
                    question_type = 'choice'
                
                # 提取选项（如果是选择题）
                options = []
                if question_type == 'choice' and re.search(r'[A-Z]、', question_text):
                    options = re.findall(r'([A-Z]、.*?)(?=[A-Z]、|$)', question_text)
                
                # 打印调试信息
                print(f"\n处理题目 {num}:")
                print(f"问题内容: {question_text}")
                print(f"题目类型: {question_type}")
                if options:
                    print(f"选项数量: {len(options)}")
                
                # 构建完整的题目信息
                full_question = {
                    'type': question_type,
                    'number': num,
                    'question': question_text,
                    'options': options,
                    'reasoning': None,
                    'answer': None,
                    'result': None,
                    'knowledge_context': None
                }
                
                # 使用Ollama分析答案
                if question_type == 'fill_in':
                    prompt = f"这是一道填空题：{question_text}\n请分析并给出答案。"
                elif question_type == 'essay':
                    prompt = f"这是一道简答题：{question_text}\n请分析并给出答案。"
                else:
                    prompt = f"这是一道选择题：{question_text}\n"
                    if options:
                        prompt += "选项：\n" + "\n".join(options) + "\n"
                    prompt += "请分析并选择正确答案。"
                
                print(f"正在分析题目 {num}...")
                analysis = self.process_with_ollama(prompt)
                if analysis:
                    full_question['reasoning'] = analysis['reasoning']
                    full_question['answer'] = analysis['answer']
                    full_question['knowledge_context'] = analysis.get('knowledge_context', '')
                    # 根据推理过程和答案判断结果
                    if "正确" in analysis['reasoning'] or "合理" in analysis['reasoning']:
                        full_question['result'] = "正确"
                    else:
                        full_question['result'] = "需要进一步验证"
                
                results.append(full_question)
                
            except Exception as e:
                print(f"处理题目时出错: {str(e)}")
                continue
        
        return results

    def generate_doc(self, results, output_path):
        """生成Word文档"""
        try:
            # 检查文件是否已存在且被占用
            if os.path.exists(output_path):
                try:
                    # 尝试重命名现有文件
                    backup_path = f"{output_path}.bak"
                    if os.path.exists(backup_path):
                        os.remove(backup_path)
                    os.rename(output_path, backup_path)
                    print(f"已备份原文件到: {backup_path}")
                except PermissionError:
                    # 如果无法重命名，尝试使用新的文件名
                    base_name, ext = os.path.splitext(output_path)
                    counter = 1
                    while True:
                        new_path = f"{base_name}_{counter}{ext}"
                        if not os.path.exists(new_path):
                            output_path = new_path
                            print(f"使用新文件名: {output_path}")
                            break
                        counter += 1
            
            doc = Document()
            
            # 添加标题
            title = doc.add_paragraph('题目分析报告')
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title.runs[0].font.size = Pt(16)
            title.runs[0].font.bold = True
            
            # 添加内容
            for item in results:
                # 添加题目编号
                q_para = doc.add_paragraph('第 ' + item["number"] + ' 题:')
                q_para.runs[0].font.bold = True
                
                # 添加题目内容
                question_para = doc.add_paragraph()
                question_para.add_run('题目：').font.bold = True
                question_para.add_run(item['question']).font.size = Pt(12)
                
                # 如果是选择题，添加选项
                if item['type'] == 'choice' and item['options']:
                    for option in item['options']:
                        option_para = doc.add_paragraph()
                        option_para.add_run(option.strip()).font.size = Pt(11)
                
                # 添加知识库内容
                if item.get('knowledge_context'):
                    knowledge_para = doc.add_paragraph()
                    knowledge_para.add_run('相关知识：').font.bold = True
                    doc.add_paragraph(item['knowledge_context']).font.size = Pt(11)
                
                # 添加答案
                answer_para = doc.add_paragraph()
                answer_para.add_run('回答：').font.bold = True
                answer_para.add_run(item['answer'] if item['answer'] else "无法确定")
                
                # 添加推理过程
                reasoning_para = doc.add_paragraph()
                reasoning_para.add_run('推理过程：').font.bold = True
                if item['reasoning']:
                    for line in item['reasoning'].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("无推理过程")
                
                # 添加结果
                result_para = doc.add_paragraph()
                result_para.add_run('结果：').font.bold = True
                result_para.add_run(item['result'] if item['result'] else "无法判断")
                
                # 添加分隔线
                doc.add_paragraph('-' * 50)
            
            # 保存文档
            doc.save(output_path)
            print(f"文档已成功保存到: {output_path}")
            
        except Exception as e:
            print(f"保存文档时发生错误: {str(e)}")
            # 尝试使用临时文件
            temp_path = "temp_analysis_results.docx"
            try:
                doc.save(temp_path)
                print(f"文档已保存到临时文件: {temp_path}")
            except Exception as e:
                print(f"无法保存到临时文件: {str(e)}")

def main():
    processor = MarkdownOllamaProcessor()
    
    # 创建输出目录
    output_dir = "md_output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"创建输出目录: {output_dir}")
    
    # 遍历 pdf_output 文件夹中的所有文件
    input_dir = "pdf_output"
    if not os.path.exists(input_dir):
        print(f"错误：输入目录 {input_dir} 不存在！")
        return
    
    # 获取所有文件
    files = [f for f in os.listdir(input_dir) if os.path.isfile(os.path.join(input_dir, f))]
    if not files:
        print(f"警告：在 {input_dir} 目录中没有找到任何文件！")
        return
    
    print(f"找到 {len(files)} 个文件需要处理")
    
    # 处理每个文件
    for file in files:
        try:
            # 构建完整的文件路径
            file_path = os.path.join(input_dir, file)
            
            print(f"\n正在处理文件: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            print(f"文件内容长度: {len(content)} 字符")
            
            questions = processor.extract_questions_and_answers(content)
            
            if not questions:
                print(f"警告：在文件 {file} 中没有找到任何题目！")
                continue
            
            # 生成输出文件名
            output_filename = f"{os.path.splitext(file)[0]}_analysis.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            # 生成文档
            processor.generate_doc(questions, output_path)
            print(f"分析报告已生成：{output_path}")
            print(f"共处理了 {len(questions)} 个题目")
            
        except Exception as e:
            print(f"处理文件 {file} 时出错: {str(e)}")
            continue

if __name__ == "__main__":
    main()